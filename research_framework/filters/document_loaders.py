"""
Document Loaders module for the Semantic Layer.

This module provides document loading capabilities for various file formats,
enabling the Knowledge Base to ingest content from PDFs, Word documents,
Excel spreadsheets, CSV files, and plain text files.

Each loader extracts text content and metadata from documents, preparing them
for embedding and storage in the vector store.

Architecture:
    - DocumentLoaderConfig: Configuration for loader behavior
    - BaseDocumentLoader: Abstract base class for all loaders
    - PDFLoader: Load PDF documents using pypdf
    - WordLoader: Load .docx files using python-docx
    - ExcelLoader: Load .xlsx files using openpyxl
    - CSVLoader: Load CSV files using pandas
    - TextLoader: Load plain text files (.txt, .md)
    - DirectoryLoader: Load all supported files from a directory
    - get_loader_for_file(): Factory function to get appropriate loader

Supported Formats:
    - PDF (.pdf): Text extraction with page metadata
    - Word (.docx): Paragraph and table extraction
    - Excel (.xlsx, .xls): Sheet and cell extraction
    - CSV (.csv): Row-by-row or chunked loading
    - Text (.txt, .md, .json, .xml): Plain text loading

Usage:
    >>> from filters.document_loaders import PDFLoader, DirectoryLoader
    >>>
    >>> # Load a single PDF
    >>> loader = PDFLoader()
    >>> docs = loader.load("report.pdf")
    >>>
    >>> # Load entire directory
    >>> dir_loader = DirectoryLoader("./documents/")
    >>> all_docs = dir_loader.load()

Example:
    >>> # Load and chunk a large document
    >>> loader = PDFLoader(chunk_size=1000, chunk_overlap=200)
    >>> docs = loader.load("large_report.pdf")
    >>> print(f"Loaded {len(docs)} chunks from PDF")
"""

import os
import re
import hashlib
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Iterator, Set
from datetime import datetime

from utils.logger import get_logger


logger = get_logger(__name__)


# =============================================================================
# Configuration
# =============================================================================

@dataclass
class DocumentLoaderConfig:
    """
    Configuration for document loaders.

    This dataclass controls how documents are loaded and processed,
    including chunking behavior and metadata extraction.

    Attributes:
        chunk_size: Maximum characters per chunk. Set to 0 or None for no chunking.
            Recommended values:
            - 500-1000: For fine-grained search
            - 1000-2000: For general use (default)
            - 2000-4000: For context-heavy retrieval
        chunk_overlap: Number of characters to overlap between chunks.
            Helps maintain context across chunk boundaries.
        include_metadata: Whether to extract and include file metadata.
        encoding: Text encoding for reading files (default: utf-8).
        skip_empty: Whether to skip empty chunks/documents.
        min_chunk_size: Minimum chunk size to keep (filters tiny chunks).
        separator_pattern: Regex pattern for preferred split points.
            Default splits on paragraphs, then sentences.

    Example:
        >>> config = DocumentLoaderConfig(
        ...     chunk_size=1000,
        ...     chunk_overlap=200,
        ...     include_metadata=True
        ... )
        >>> loader = PDFLoader(config=config)
    """

    chunk_size: int = 1500
    chunk_overlap: int = 200
    include_metadata: bool = True
    encoding: str = "utf-8"
    skip_empty: bool = True
    min_chunk_size: int = 50
    separator_pattern: str = r"\n\n|\n|\.(?=\s)"  # Paragraphs, newlines, sentences


# =============================================================================
# Document Data Structure
# =============================================================================

@dataclass
class LoadedDocument:
    """
    Represents a loaded document or document chunk.

    This dataclass holds the extracted content and metadata from a loaded
    document. When chunking is enabled, each chunk becomes a separate
    LoadedDocument with its own metadata.

    Attributes:
        id: Unique identifier for the document/chunk.
        content: The text content.
        metadata: Dictionary of metadata (source, page, chunk_index, etc.).

    Example:
        >>> doc = LoadedDocument(
        ...     id="report_p1_c0",
        ...     content="AI is transforming healthcare...",
        ...     metadata={"source": "report.pdf", "page": 1, "chunk_index": 0}
        ... )
    """

    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for vector store ingestion."""
        return {
            "id": self.id,
            "content": self.content,
            "metadata": self.metadata,
        }


# =============================================================================
# Text Chunking Utility
# =============================================================================

class TextChunker:
    """
    Utility class for splitting text into chunks.

    This class implements intelligent text chunking that respects natural
    text boundaries (paragraphs, sentences) while maintaining a target
    chunk size with overlap.

    The chunking algorithm:
    1. Splits text on preferred separators (paragraphs, sentences)
    2. Combines splits into chunks up to max_size
    3. Adds overlap from previous chunk for context continuity

    Attributes:
        max_size: Maximum chunk size in characters.
        overlap: Number of overlapping characters between chunks.
        separator_pattern: Regex for preferred split points.

    Example:
        >>> chunker = TextChunker(max_size=1000, overlap=200)
        >>> chunks = chunker.chunk("Long text content here...")
        >>> print(f"Split into {len(chunks)} chunks")
    """

    def __init__(
        self,
        max_size: int = 1500,
        overlap: int = 200,
        separator_pattern: str = r"\n\n|\n|\.(?=\s)",
    ):
        """
        Initialize the text chunker.

        Args:
            max_size: Maximum characters per chunk.
            overlap: Overlap between consecutive chunks.
            separator_pattern: Regex pattern for split points.
        """
        self.max_size = max_size
        self.overlap = overlap
        self.separator_pattern = separator_pattern

    def chunk(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: The text to chunk.

        Returns:
            List[str]: List of text chunks.

        Example:
            >>> chunks = chunker.chunk("Paragraph 1.\\n\\nParagraph 2.\\n\\nParagraph 3.")
        """
        if not text or len(text) <= self.max_size:
            return [text] if text else []

        # Split on separators while keeping the separators
        splits = re.split(f"({self.separator_pattern})", text)

        # Recombine splits with their separators
        segments = []
        current = ""
        for i, split in enumerate(splits):
            # Check if this is a separator (odd indices after re.split with groups)
            if i % 2 == 1:
                current += split
            else:
                if current:
                    segments.append(current)
                current = split
        if current:
            segments.append(current)

        # Build chunks from segments
        chunks = []
        current_chunk = ""

        for segment in segments:
            # If adding this segment exceeds max_size
            if len(current_chunk) + len(segment) > self.max_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())

                    # Start new chunk with overlap from previous
                    if self.overlap > 0 and len(current_chunk) > self.overlap:
                        # Find a good split point for overlap
                        overlap_text = current_chunk[-self.overlap:]
                        # Try to start overlap at a word boundary
                        space_idx = overlap_text.find(" ")
                        if space_idx > 0:
                            overlap_text = overlap_text[space_idx + 1:]
                        current_chunk = overlap_text + segment
                    else:
                        current_chunk = segment
                else:
                    # Segment itself is larger than max_size, force split
                    for i in range(0, len(segment), self.max_size - self.overlap):
                        chunk = segment[i:i + self.max_size]
                        if chunk.strip():
                            chunks.append(chunk.strip())
                    current_chunk = ""
            else:
                current_chunk += segment

        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks


# =============================================================================
# Base Document Loader (Abstract)
# =============================================================================

class BaseDocumentLoader(ABC):
    """
    Abstract base class for document loaders.

    This class defines the interface that all document loaders must implement,
    providing a consistent API for loading documents regardless of file format.

    Subclasses must implement:
        - load(): Load and return documents from the source
        - supported_extensions: Property returning supported file extensions

    Attributes:
        config: Loader configuration settings.
        chunker: TextChunker instance for splitting content.
    """

    def __init__(self, config: Optional[DocumentLoaderConfig] = None):
        """
        Initialize the document loader.

        Args:
            config: Loader configuration. Uses defaults if None.
        """
        self.config = config or DocumentLoaderConfig()
        self.chunker = TextChunker(
            max_size=self.config.chunk_size,
            overlap=self.config.chunk_overlap,
            separator_pattern=self.config.separator_pattern,
        )

    @property
    @abstractmethod
    def supported_extensions(self) -> Set[str]:
        """Return set of supported file extensions (e.g., {'.pdf', '.PDF'})."""
        pass

    @abstractmethod
    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load documents from the source.

        Args:
            source: File path or other source identifier.

        Returns:
            List[LoadedDocument]: Loaded document(s) or chunks.
        """
        pass

    def _generate_id(self, source: str, *args) -> str:
        """
        Generate a unique document ID.

        Args:
            source: Source file path or identifier.
            *args: Additional components for the ID (page number, chunk index).

        Returns:
            str: Unique document ID.
        """
        components = [str(source)] + [str(a) for a in args]
        id_string = "_".join(components)
        # Use hash for long paths
        if len(id_string) > 100:
            hash_suffix = hashlib.md5(id_string.encode()).hexdigest()[:8]
            return f"{Path(source).stem}_{hash_suffix}"
        return id_string.replace("/", "_").replace("\\", "_").replace(" ", "_")

    def _get_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from a file.

        Args:
            file_path: Path to the file.

        Returns:
            Dict: File metadata including name, size, modified time.
        """
        if not self.config.include_metadata:
            return {}

        try:
            stat = file_path.stat()
            return {
                "source": str(file_path),
                "filename": file_path.name,
                "file_size": stat.st_size,
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "file_type": file_path.suffix.lower(),
            }
        except OSError:
            return {"source": str(file_path)}

    def _chunk_content(
        self,
        content: str,
        source: str,
        base_metadata: Dict[str, Any],
    ) -> List[LoadedDocument]:
        """
        Chunk content into multiple documents if chunking is enabled.

        Args:
            content: Text content to chunk.
            source: Source identifier for ID generation.
            base_metadata: Metadata to include in all chunks.

        Returns:
            List[LoadedDocument]: List of document chunks.
        """
        # Skip empty content
        if not content or (self.config.skip_empty and not content.strip()):
            return []

        # No chunking if disabled or content is small enough
        if self.config.chunk_size <= 0 or len(content) <= self.config.chunk_size:
            return [
                LoadedDocument(
                    id=self._generate_id(source),
                    content=content,
                    metadata=base_metadata,
                )
            ]

        # Chunk the content
        chunks = self.chunker.chunk(content)
        documents = []

        for i, chunk in enumerate(chunks):
            # Skip chunks that are too small
            if len(chunk) < self.config.min_chunk_size:
                continue

            chunk_metadata = {
                **base_metadata,
                "chunk_index": i,
                "total_chunks": len(chunks),
            }

            documents.append(
                LoadedDocument(
                    id=self._generate_id(source, f"chunk{i}"),
                    content=chunk,
                    metadata=chunk_metadata,
                )
            )

        return documents


# =============================================================================
# PDF Loader
# =============================================================================

class PDFLoader(BaseDocumentLoader):
    """
    Load documents from PDF files.

    This loader uses pypdf to extract text from PDF files. It supports
    page-by-page extraction with optional chunking within pages.

    Features:
        - Page-level metadata (page numbers)
        - Text extraction from all pages
        - Handles multi-column layouts reasonably well
        - Chunking support for large pages

    Example:
        >>> loader = PDFLoader()
        >>> docs = loader.load("report.pdf")
        >>> for doc in docs:
        ...     print(f"Page {doc.metadata.get('page')}: {doc.content[:100]}...")
    """

    @property
    def supported_extensions(self) -> Set[str]:
        """Supported PDF extensions."""
        return {".pdf", ".PDF"}

    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load a PDF file.

        Args:
            source: Path to the PDF file.

        Returns:
            List[LoadedDocument]: Extracted documents/chunks.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ImportError: If pypdf is not installed.
        """
        file_path = Path(source)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF loading. "
                "Install with: pip install pypdf"
            )

        logger.info(f"Loading PDF: {file_path}")
        documents = []
        base_metadata = self._get_file_metadata(file_path)

        try:
            reader = PdfReader(str(file_path))
            total_pages = len(reader.pages)
            base_metadata["total_pages"] = total_pages

            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text() or ""

                    if not text.strip() and self.config.skip_empty:
                        continue

                    page_metadata = {
                        **base_metadata,
                        "page": page_num,
                    }

                    # Chunk page content
                    page_docs = self._chunk_content(
                        content=text,
                        source=f"{file_path}_p{page_num}",
                        base_metadata=page_metadata,
                    )
                    documents.extend(page_docs)

                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")

            logger.info(f"Loaded {len(documents)} chunks from PDF ({total_pages} pages)")

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise

        return documents


# =============================================================================
# Word Document Loader
# =============================================================================

class WordLoader(BaseDocumentLoader):
    """
    Load documents from Microsoft Word (.docx) files.

    This loader uses python-docx to extract text from Word documents,
    including paragraphs and tables.

    Features:
        - Paragraph extraction
        - Table extraction with row/cell structure
        - Style information in metadata (optional)

    Example:
        >>> loader = WordLoader()
        >>> docs = loader.load("document.docx")
    """

    @property
    def supported_extensions(self) -> Set[str]:
        """Supported Word extensions."""
        return {".docx", ".DOCX"}

    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load a Word document.

        Args:
            source: Path to the .docx file.

        Returns:
            List[LoadedDocument]: Extracted documents/chunks.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ImportError: If python-docx is not installed.
        """
        file_path = Path(source)
        if not file_path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for Word loading. "
                "Install with: pip install python-docx"
            )

        logger.info(f"Loading Word document: {file_path}")
        base_metadata = self._get_file_metadata(file_path)

        try:
            doc = Document(str(file_path))

            # Extract all text
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))

            full_text = "\n\n".join(text_parts)

            base_metadata["paragraph_count"] = len(doc.paragraphs)
            base_metadata["table_count"] = len(doc.tables)

            documents = self._chunk_content(
                content=full_text,
                source=str(file_path),
                base_metadata=base_metadata,
            )

            logger.info(f"Loaded {len(documents)} chunks from Word document")
            return documents

        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            raise


# =============================================================================
# Excel Loader
# =============================================================================

class ExcelLoader(BaseDocumentLoader):
    """
    Load documents from Excel (.xlsx, .xls) files.

    This loader uses openpyxl for .xlsx files and can optionally use
    pandas for more complex data handling. Each sheet can be loaded
    as a separate document or combined.

    Features:
        - Multi-sheet support
        - Row-by-row or sheet-level loading
        - Header detection
        - Cell value extraction

    Example:
        >>> loader = ExcelLoader()
        >>> docs = loader.load("data.xlsx")
    """

    def __init__(
        self,
        config: Optional[DocumentLoaderConfig] = None,
        combine_sheets: bool = True,
        include_headers: bool = True,
    ):
        """
        Initialize Excel loader.

        Args:
            config: Loader configuration.
            combine_sheets: Whether to combine all sheets into one document.
            include_headers: Whether to include column headers in output.
        """
        super().__init__(config)
        self.combine_sheets = combine_sheets
        self.include_headers = include_headers

    @property
    def supported_extensions(self) -> Set[str]:
        """Supported Excel extensions."""
        return {".xlsx", ".xls", ".XLSX", ".XLS"}

    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load an Excel file.

        Args:
            source: Path to the Excel file.

        Returns:
            List[LoadedDocument]: Extracted documents/chunks.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ImportError: If openpyxl is not installed.
        """
        file_path = Path(source)
        if not file_path.exists():
            raise FileNotFoundError(f"Excel file not found: {file_path}")

        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError(
                "openpyxl is required for Excel loading. "
                "Install with: pip install openpyxl"
            )

        logger.info(f"Loading Excel file: {file_path}")
        base_metadata = self._get_file_metadata(file_path)
        documents = []

        try:
            wb = load_workbook(str(file_path), data_only=True)
            base_metadata["sheet_count"] = len(wb.sheetnames)

            all_text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = []
                headers = []

                for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                    # Filter out completely empty rows
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if not any(row_values):
                        continue

                    # First non-empty row as headers
                    if row_idx == 0 and self.include_headers:
                        headers = row_values
                        sheet_text.append(" | ".join(headers))
                    else:
                        # Format row with or without headers
                        if headers and self.include_headers:
                            row_str = ", ".join(
                                f"{h}: {v}" for h, v in zip(headers, row_values) if v
                            )
                        else:
                            row_str = " | ".join(v for v in row_values if v)
                        if row_str:
                            sheet_text.append(row_str)

                if sheet_text:
                    sheet_content = f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_text)
                    all_text.append(sheet_content)

                    if not self.combine_sheets:
                        sheet_metadata = {
                            **base_metadata,
                            "sheet_name": sheet_name,
                        }
                        sheet_docs = self._chunk_content(
                            content=sheet_content,
                            source=f"{file_path}_{sheet_name}",
                            base_metadata=sheet_metadata,
                        )
                        documents.extend(sheet_docs)

            if self.combine_sheets and all_text:
                full_text = "\n\n".join(all_text)
                documents = self._chunk_content(
                    content=full_text,
                    source=str(file_path),
                    base_metadata=base_metadata,
                )

            logger.info(f"Loaded {len(documents)} chunks from Excel file")
            return documents

        except Exception as e:
            logger.error(f"Error loading Excel file {file_path}: {e}")
            raise


# =============================================================================
# CSV Loader
# =============================================================================

class CSVLoader(BaseDocumentLoader):
    """
    Load documents from CSV files.

    This loader uses pandas to read CSV files and converts rows into
    text documents. It supports various CSV configurations and can
    load rows individually or in groups.

    Features:
        - Configurable delimiter and encoding
        - Header detection
        - Row grouping for chunking
        - Handles large files efficiently

    Example:
        >>> loader = CSVLoader()
        >>> docs = loader.load("data.csv")
    """

    def __init__(
        self,
        config: Optional[DocumentLoaderConfig] = None,
        delimiter: str = ",",
        rows_per_chunk: int = 50,
    ):
        """
        Initialize CSV loader.

        Args:
            config: Loader configuration.
            delimiter: CSV field delimiter.
            rows_per_chunk: Number of rows to combine per document chunk.
        """
        super().__init__(config)
        self.delimiter = delimiter
        self.rows_per_chunk = rows_per_chunk

    @property
    def supported_extensions(self) -> Set[str]:
        """Supported CSV extensions."""
        return {".csv", ".CSV", ".tsv", ".TSV"}

    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load a CSV file.

        Args:
            source: Path to the CSV file.

        Returns:
            List[LoadedDocument]: Extracted documents/chunks.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ImportError: If pandas is not installed.
        """
        file_path = Path(source)
        if not file_path.exists():
            raise FileNotFoundError(f"CSV file not found: {file_path}")

        try:
            import pandas as pd
        except ImportError:
            raise ImportError(
                "pandas is required for CSV loading. "
                "Install with: pip install pandas"
            )

        logger.info(f"Loading CSV file: {file_path}")
        base_metadata = self._get_file_metadata(file_path)
        documents = []

        try:
            # Detect delimiter for TSV files
            delimiter = "\t" if file_path.suffix.lower() in [".tsv"] else self.delimiter

            df = pd.read_csv(
                file_path,
                delimiter=delimiter,
                encoding=self.config.encoding,
            )

            base_metadata["row_count"] = len(df)
            base_metadata["column_count"] = len(df.columns)
            base_metadata["columns"] = list(df.columns)

            # Convert rows to text
            headers = list(df.columns)

            for chunk_idx in range(0, len(df), self.rows_per_chunk):
                chunk_df = df.iloc[chunk_idx:chunk_idx + self.rows_per_chunk]
                rows_text = []

                for _, row in chunk_df.iterrows():
                    row_str = ", ".join(
                        f"{col}: {val}"
                        for col, val in row.items()
                        if pd.notna(val) and str(val).strip()
                    )
                    if row_str:
                        rows_text.append(row_str)

                if rows_text:
                    chunk_metadata = {
                        **base_metadata,
                        "start_row": chunk_idx,
                        "end_row": min(chunk_idx + self.rows_per_chunk, len(df)),
                    }

                    documents.append(
                        LoadedDocument(
                            id=self._generate_id(str(file_path), f"rows{chunk_idx}"),
                            content="\n".join(rows_text),
                            metadata=chunk_metadata,
                        )
                    )

            logger.info(f"Loaded {len(documents)} chunks from CSV ({len(df)} rows)")
            return documents

        except Exception as e:
            logger.error(f"Error loading CSV file {file_path}: {e}")
            raise


# =============================================================================
# Text File Loader
# =============================================================================

class TextLoader(BaseDocumentLoader):
    """
    Load documents from plain text files.

    This loader handles text files including .txt, .md, .json, .xml,
    and other text-based formats.

    Features:
        - Multiple encoding support
        - Markdown and JSON awareness
        - Simple and fast loading

    Example:
        >>> loader = TextLoader()
        >>> docs = loader.load("readme.md")
    """

    @property
    def supported_extensions(self) -> Set[str]:
        """Supported text extensions."""
        return {
            ".txt", ".TXT",
            ".md", ".MD", ".markdown",
            ".json", ".JSON",
            ".xml", ".XML",
            ".yaml", ".yml", ".YAML", ".YML",
            ".log", ".LOG",
            ".rst", ".RST",
        }

    def load(self, source: Union[str, Path]) -> List[LoadedDocument]:
        """
        Load a text file.

        Args:
            source: Path to the text file.

        Returns:
            List[LoadedDocument]: Extracted documents/chunks.

        Raises:
            FileNotFoundError: If the file doesn't exist.
        """
        file_path = Path(source)
        if not file_path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")

        logger.info(f"Loading text file: {file_path}")
        base_metadata = self._get_file_metadata(file_path)

        try:
            with open(file_path, "r", encoding=self.config.encoding) as f:
                content = f.read()

            documents = self._chunk_content(
                content=content,
                source=str(file_path),
                base_metadata=base_metadata,
            )

            logger.info(f"Loaded {len(documents)} chunks from text file")
            return documents

        except UnicodeDecodeError:
            # Try with different encoding
            logger.warning(f"UTF-8 decode failed, trying latin-1 for {file_path}")
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()

            documents = self._chunk_content(
                content=content,
                source=str(file_path),
                base_metadata=base_metadata,
            )
            return documents

        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise


# =============================================================================
# Directory Loader
# =============================================================================

class DirectoryLoader:
    """
    Load all supported documents from a directory.

    This loader automatically detects file types and uses the appropriate
    loader for each file. It supports recursive directory traversal and
    file filtering.

    Features:
        - Automatic file type detection
        - Recursive directory scanning
        - File pattern filtering
        - Progress tracking for large directories

    Attributes:
        directory: Path to the directory to load.
        recursive: Whether to scan subdirectories.
        file_patterns: Optional list of glob patterns to filter files.
        config: Shared configuration for all loaders.

    Example:
        >>> loader = DirectoryLoader("./documents/", recursive=True)
        >>> docs = loader.load()
        >>> print(f"Loaded {len(docs)} documents from directory")
    """

    # Mapping of extensions to loader classes
    LOADER_MAP = {
        ".pdf": PDFLoader,
        ".docx": WordLoader,
        ".xlsx": ExcelLoader,
        ".xls": ExcelLoader,
        ".csv": CSVLoader,
        ".tsv": CSVLoader,
        ".txt": TextLoader,
        ".md": TextLoader,
        ".json": TextLoader,
        ".xml": TextLoader,
        ".yaml": TextLoader,
        ".yml": TextLoader,
    }

    def __init__(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        file_patterns: Optional[List[str]] = None,
        config: Optional[DocumentLoaderConfig] = None,
    ):
        """
        Initialize directory loader.

        Args:
            directory: Path to the directory.
            recursive: Whether to scan subdirectories.
            file_patterns: Optional glob patterns (e.g., ["*.pdf", "*.docx"]).
            config: Shared configuration for loaders.
        """
        self.directory = Path(directory)
        self.recursive = recursive
        self.file_patterns = file_patterns
        self.config = config or DocumentLoaderConfig()

        # Cache loader instances
        self._loaders: Dict[str, BaseDocumentLoader] = {}

    def _get_loader(self, extension: str) -> Optional[BaseDocumentLoader]:
        """
        Get or create a loader for the given extension.

        Args:
            extension: File extension (e.g., ".pdf").

        Returns:
            BaseDocumentLoader or None if extension not supported.
        """
        ext_lower = extension.lower()
        if ext_lower not in self._loaders:
            loader_class = self.LOADER_MAP.get(ext_lower)
            if loader_class:
                self._loaders[ext_lower] = loader_class(config=self.config)
        return self._loaders.get(ext_lower)

    def _get_files(self) -> List[Path]:
        """
        Get list of files to load.

        Returns:
            List[Path]: Paths to files matching criteria.
        """
        if not self.directory.exists():
            raise FileNotFoundError(f"Directory not found: {self.directory}")

        files = []

        # Use glob patterns if specified
        if self.file_patterns:
            for pattern in self.file_patterns:
                if self.recursive:
                    files.extend(self.directory.rglob(pattern))
                else:
                    files.extend(self.directory.glob(pattern))
        else:
            # Get all files with supported extensions
            pattern = "**/*" if self.recursive else "*"
            for path in self.directory.glob(pattern):
                if path.is_file() and path.suffix.lower() in self.LOADER_MAP:
                    files.append(path)

        return sorted(set(files))

    def load(self) -> List[LoadedDocument]:
        """
        Load all documents from the directory.

        Returns:
            List[LoadedDocument]: All loaded documents/chunks.

        Example:
            >>> loader = DirectoryLoader("./docs/")
            >>> all_docs = loader.load()
        """
        logger.info(f"Loading documents from directory: {self.directory}")
        files = self._get_files()
        logger.info(f"Found {len(files)} files to load")

        all_documents = []
        errors = []

        for file_path in files:
            loader = self._get_loader(file_path.suffix)
            if loader is None:
                logger.warning(f"No loader for extension: {file_path.suffix}")
                continue

            try:
                docs = loader.load(file_path)
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                errors.append((file_path, str(e)))

        logger.info(
            f"Loaded {len(all_documents)} documents from {len(files)} files "
            f"({len(errors)} errors)"
        )

        return all_documents

    def load_iter(self) -> Iterator[LoadedDocument]:
        """
        Iterate over documents one at a time (memory efficient).

        Yields:
            LoadedDocument: Documents one at a time.

        Example:
            >>> for doc in loader.load_iter():
            ...     process(doc)
        """
        files = self._get_files()

        for file_path in files:
            loader = self._get_loader(file_path.suffix)
            if loader is None:
                continue

            try:
                docs = loader.load(file_path)
                yield from docs
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")


# =============================================================================
# Factory Function
# =============================================================================

def get_loader_for_file(
    file_path: Union[str, Path],
    config: Optional[DocumentLoaderConfig] = None,
) -> Optional[BaseDocumentLoader]:
    """
    Get the appropriate loader for a file based on its extension.

    This factory function examines the file extension and returns an
    instance of the appropriate loader class.

    Args:
        file_path: Path to the file.
        config: Optional loader configuration.

    Returns:
        BaseDocumentLoader or None if format not supported.

    Example:
        >>> loader = get_loader_for_file("report.pdf")
        >>> if loader:
        ...     docs = loader.load("report.pdf")
    """
    path = Path(file_path)
    ext_lower = path.suffix.lower()

    loader_class = DirectoryLoader.LOADER_MAP.get(ext_lower)
    if loader_class:
        return loader_class(config=config)

    return None


def load_document(
    file_path: Union[str, Path],
    config: Optional[DocumentLoaderConfig] = None,
) -> List[LoadedDocument]:
    """
    Convenience function to load a single document.

    This function automatically selects the appropriate loader and
    loads the document.

    Args:
        file_path: Path to the document.
        config: Optional loader configuration.

    Returns:
        List[LoadedDocument]: Loaded documents/chunks.

    Raises:
        ValueError: If file format is not supported.
        FileNotFoundError: If file doesn't exist.

    Example:
        >>> docs = load_document("report.pdf")
        >>> print(f"Loaded {len(docs)} chunks")
    """
    loader = get_loader_for_file(file_path, config)
    if loader is None:
        raise ValueError(f"Unsupported file format: {Path(file_path).suffix}")

    return loader.load(file_path)
